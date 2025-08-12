from __future__ import annotations

import pathlib
from pathlib import Path
from typing import List, Tuple, Dict
from docx import Document
from docx.oxml.ns import qn
import logging
import hashlib

from scripts.utils.image_utils import (
    infer_project_root,
    get_project_image_dir,
    save_image_blob,
    generate_image_filename,
)

logger = logging.getLogger("docx_ingestor")


def load_docx(path: str | pathlib.Path) -> List[Tuple[str, dict]]:
    """Extract text and image references from a .docx file as
    (text, metadata) chunks."""
    if not isinstance(path, Path):
        path = Path(path)

    document = Document(path)
    project_root = infer_project_root(path)
    image_dir = get_project_image_dir(project_root.name)
    doc_id = str(path)

    print(f"[loader] Writing image to: {image_dir}")

    segments: List[Tuple[str, dict]] = []

    # Track saved images to avoid duplicates
    saved_images: Dict[str, str] = {}  # rId -> saved_path
    img_counter = 0  # Global counter for image numbering

    for para_idx, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        meta = {
            "doc_type": "docx",
            "paragraph_number": para_idx,
            "source_filepath": str(path),
            "doc_id": doc_id,
        }

        image_paths = []

        # Look for blips in this paragraph
        blips = paragraph._element.xpath(".//a:blip")

        for blip in blips:
            rId = blip.get(qn("r:embed"))
            if not rId:
                continue

            # Check if we already saved this image
            if rId in saved_images:
                # Reuse the existing path
                image_paths.append(saved_images[rId])
                print(
                    f"[DEBUG] Paragraph {para_idx} → reusing image {rId}: "
                    f"{saved_images[rId]}"
                )
            else:
                # Get the image part
                image_part = document.part.related_parts.get(rId)
                if not image_part:
                    continue

                # Save the new image
                img_name = generate_image_filename(
                    doc_id=doc_id,
                    page_number=para_idx,
                    img_index=img_counter,
                )
                img_path = image_dir / img_name
                save_image_blob(image_part.blob, img_path)

                try:
                    rel_path = img_path.resolve().relative_to(
                        (project_root / "input").resolve()
                    )
                    saved_path = str(rel_path)
                except ValueError:
                    saved_path = f"cache/images/{img_path.name}"

                saved_images[rId] = saved_path
                image_paths.append(saved_path)
                img_counter += 1

                print(
                    f"[DEBUG] Paragraph {para_idx} → saved new image {rId}: "
                    f"{saved_path}"
                )
                logger.info(
                    f"[DOCX] Paragraph {para_idx} → saved image {rId}: {saved_path}"
                )

        # Add image paths to metadata if any were found
        if image_paths:
            meta["image_paths"] = image_paths
            print(
                f"[DEBUG] Paragraph {para_idx} → extracted {len(image_paths)} "
                f"image(s): {image_paths}"
            )
            logger.info(
                f"[DOCX] [DEBUG] Paragraph {para_idx} → extracted {len(image_paths)} "
                f"image(s): {image_paths}"
            )

        # Create segment if there's text or images
        if text or image_paths:
            segments.append((text or "[Image-only content]", meta))
        else:
            print(
                f"[DEBUG] Paragraph {para_idx} was skipped — "
                f"no text and no images recorded."
            )
            logger.debug(
                f"[DOCX] [DEBUG] Paragraph {para_idx} was skipped — "
                f"no text and no images recorded."
            )

    # Debug: Print all segments with images
    print(f"\n[DEBUG] All segments with images:")
    for idx, (text, meta) in enumerate(segments):
        if 'image_paths' in meta:
            print(
                f"  Segment {idx}: Paragraph {meta['paragraph_number']}, "
                f"Images: {meta['image_paths']}"
            )

    print(
        f"[INFO] Extracted {sum('image_paths' in m for _, m in segments)} "
        f"image-attached chunks from {path.name}"
    )
    print(f"[INFO] Total segments: {len(segments)}")

    # Add tables
    for tbl_idx, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                rows.append(" | ".join(row_cells))
        if rows:
            tbl_text = "\n".join(rows)
            meta = {
                "doc_type": "docx",
                "table_number": tbl_idx,
                "source_filepath": str(path),
                "doc_id": doc_id,
            }
            segments.append((tbl_text, meta))

    # Final verification
    print(f"\n[FINAL DEBUG] Returning {len(segments)} segments")
    for i, (text, meta) in enumerate(segments):
        if 'image_paths' in meta:
            print(
                f"  Segment {i}: {meta.get('doc_type')} para "
                f"{meta.get('paragraph_number', 'N/A')}, images: {meta['image_paths']}"
            )

    return segments
