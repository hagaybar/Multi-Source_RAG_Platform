from __future__ import annotations

import pathlib
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.oxml.ns import qn


def _infer_project_root(file_path: Path) -> Path:
    """Best-effort guess of the project root from a raw file path."""
    parts = file_path.resolve().parts
    if "projects" in parts:
        idx = parts.index("projects")
        if idx + 1 < len(parts):
            return Path(*parts[: idx + 2])
    return file_path.parent


def _save_image(blob: bytes, project_root: Path, filename: str) -> str:
    rel_dir = Path("input") / "cache" / "images"
    out_path = project_root / rel_dir / filename
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "wb") as f:
        f.write(blob)
    return str(rel_dir / filename)


def load_docx(path: str | pathlib.Path) -> List[Tuple[str, dict]]:
    """Extract text and images from a .docx file.

    Returns a list of ``(text, metadata)`` tuples where ``metadata`` may contain
    ``image_path`` if an image was found in that paragraph.
    """
    if not isinstance(path, Path):
        path = Path(path)

    document = Document(path)
    project_root = _infer_project_root(path)
    segments: List[Tuple[str, dict]] = []
    file_stem = path.stem

    for para_idx, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue

        base_meta = {"doc_type": "docx", "paragraph_number": para_idx}
        img_count = 0
        for run in paragraph.runs:
            blips = run._element.xpath(".//a:blip")
            for blip in blips:
                rId = blip.get(qn("r:embed"))
                image_part = document.part.related_parts.get(rId)
                if image_part is None:
                    continue
                img_count += 1
                img_name = f"{file_stem}_page{para_idx}_img{img_count}.png"
                rel_path = _save_image(image_part.blob, project_root, img_name)
                meta = base_meta.copy()
                meta["image_path"] = rel_path
                segments.append((text, meta))

        if img_count == 0:
            segments.append((text, base_meta))

    # Tables are appended as additional segments
    for tbl_idx, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                rows.append(" | ".join(row_cells))
        if rows:
            tbl_text = "\n".join(rows)
            meta = {
                "doc_type": "docx",
                "table_number": tbl_idx,
            }
            segments.append((tbl_text, meta))

    return segments
